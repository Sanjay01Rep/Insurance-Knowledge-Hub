"""Build downloadable copies of an answer. Do not include API keys."""

from __future__ import annotations

from io import BytesIO


def answer_text(answer: str, sources: list[dict], model_note: str = "") -> str:
    lines = ["Insurance Learning Assistant", ""]
    if model_note:
        lines.append(model_note)
        lines.append("")
    lines.append((answer or "").strip())
    doc_sources = [s for s in sources if (s.get("kind") or "document") == "document"]
    web_sources = [s for s in sources if s.get("kind") == "web"]
    if doc_sources:
        lines.extend(["", "Sources (your documents):"])
        for source in doc_sources:
            lines.append(f"- {source.get('label') or ''}")
    if web_sources:
        lines.extend(["", "From the web (not from your documents):"])
        for source in web_sources:
            label = source.get("label") or ""
            url = source.get("url") or ""
            lines.append(f"- {label}" + (f" — {url}" if url else ""))
    return "\n".join(lines).strip() + "\n"


def answer_docx_bytes(answer: str, sources: list[dict], model_note: str = "") -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Insurance Learning Assistant", level=0)
    if model_note:
        doc.add_paragraph(model_note)
    for paragraph in (answer or "").split("\n"):
        doc.add_paragraph(paragraph)
    doc_sources = [s for s in sources if (s.get("kind") or "document") == "document"]
    web_sources = [s for s in sources if s.get("kind") == "web"]
    if doc_sources:
        doc.add_heading("Sources (your documents)", level=1)
        for source in doc_sources:
            doc.add_paragraph(source.get("label") or "", style="List Bullet")
    if web_sources:
        doc.add_heading("From the web (not from your documents)", level=1)
        for source in web_sources:
            label = source.get("label") or ""
            url = source.get("url") or ""
            doc.add_paragraph(label + (f" — {url}" if url else ""), style="List Bullet")
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def answer_pdf_bytes(answer: str, sources: list[dict], model_note: str = "") -> bytes:
    text = answer_text(answer, sources, model_note)
    lines = _wrap_pdf_lines(text, 90)
    content_lines = []
    y = 800
    for line in lines[:60]:
        escaped = (
            line.replace("\\", "\\\\")
            .replace("(", "\\(")
            .replace(")", "\\)")
        )
        content_lines.append(f"BT /F1 11 Tf 50 {y} Td ({escaped}) Tj ET")
        y -= 16
        if y < 50:
            break
    stream = "\n".join(content_lines).encode("latin-1", errors="replace")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length "
        + str(len(stream)).encode("ascii")
        + b" >>\nstream\n"
        + stream
        + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out.extend(f"{i} 0 obj\n".encode("ascii"))
        out.extend(obj)
        out.extend(b"\nendobj\n")
    xref = len(out)
    out.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    out.extend(b"0000000000 65535 f \n")
    for off in offsets[1:]:
        out.extend(f"{off:010d} 00000 n \n".encode("ascii"))
    out.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode(
            "ascii"
        )
    )
    return bytes(out)


def _wrap_pdf_lines(text: str, width: int) -> list[str]:
    lines: list[str] = []
    for raw in (text or "").replace("\r", "").split("\n"):
        line = raw.strip() if not raw.strip() else raw
        while len(line) > width:
            cut = line.rfind(" ", 0, width)
            if cut < 20:
                cut = width
            lines.append(line[:cut])
            line = line[cut:].lstrip()
        lines.append(line)
    return lines
